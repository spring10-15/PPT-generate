from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
MARKDOWN_PATH = ROOT / "docs" / "PRD.md"
OUTPUT_DIR = ROOT / "output" / "doc"
TMP_DIR = ROOT / "tmp" / "docs"
OUTPUT_PATH = OUTPUT_DIR / "固定模板PPT生成Demo-PRD.docx"
FLOW_OVERVIEW_PATH = TMP_DIR / "prd-flow-overview.png"
FLOW_EXPORT_PATH = TMP_DIR / "prd-flow-export.png"

FONT_CANDIDATES = [
    "/System/Library/Fonts/PingFang.ttc",
    "/System/Library/Fonts/Supplemental/PingFang.ttc",
    "/System/Library/Fonts/Supplemental/Songti.ttc",
    "/System/Library/Fonts/Hiragino Sans GB.ttc",
]


@dataclass
class FlowNode:
    text: str
    x: int
    y: int
    w: int = 220
    h: int = 72
    fill: str = "#F4F8FB"
    outline: str = "#0070BF"


def find_font(size: int) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    for candidate in FONT_CANDIDATES:
        path = Path(candidate)
        if path.exists():
            return ImageFont.truetype(str(path), size=size)
    return ImageFont.load_default()


def draw_centered_text(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str, font) -> None:
    left, top, right, bottom = box
    max_width = right - left - 20
    lines: list[str] = []
    current = ""

    for char in text:
        probe = f"{current}{char}"
        width = draw.textbbox((0, 0), probe, font=font)[2]
        if width <= max_width or not current:
            current = probe
        else:
            lines.append(current)
            current = char
    if current:
        lines.append(current)

    line_height = draw.textbbox((0, 0), "示例", font=font)[3] + 4
    total_height = line_height * len(lines)
    y = top + ((bottom - top) - total_height) / 2

    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font)
        x = left + ((right - left) - (bbox[2] - bbox[0])) / 2
        draw.text((x, y), line, fill="#17324D", font=font)
        y += line_height


def draw_arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], color: str = "#0070BF") -> None:
    draw.line([start, end], fill=color, width=4)
    if start[0] == end[0]:
        direction = 1 if end[1] > start[1] else -1
        arrow = [
            (end[0], end[1]),
            (end[0] - 10, end[1] - 14 * direction),
            (end[0] + 10, end[1] - 14 * direction),
        ]
    else:
        direction = 1 if end[0] > start[0] else -1
        arrow = [
            (end[0], end[1]),
            (end[0] - 14 * direction, end[1] - 10),
            (end[0] - 14 * direction, end[1] + 10),
        ]
    draw.polygon(arrow, fill=color)


def create_flowchart(path: Path, title: str, nodes: list[FlowNode], arrows: Iterable[tuple[int, int]]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    canvas = Image.new("RGB", (1320, 880), "#FFFFFF")
    draw = ImageDraw.Draw(canvas)
    title_font = find_font(32)
    node_font = find_font(24)

    draw.rounded_rectangle((24, 24, 1296, 856), radius=28, outline="#D7E2EB", width=2, fill="#FFFFFF")
    draw.text((48, 42), title, fill="#0070BF", font=title_font)

    for node in nodes:
        rect = (node.x, node.y, node.x + node.w, node.y + node.h)
        draw.rounded_rectangle(rect, radius=18, fill=node.fill, outline=node.outline, width=3)
        draw_centered_text(draw, rect, node.text, node_font)

    for start_idx, end_idx in arrows:
        start_node = nodes[start_idx]
        end_node = nodes[end_idx]
        if start_node.x == end_node.x:
            start = (start_node.x + start_node.w // 2, start_node.y + start_node.h)
            end = (end_node.x + end_node.w // 2, end_node.y)
        else:
            start = (start_node.x + start_node.w, start_node.y + start_node.h // 2)
            end = (end_node.x, end_node.y + end_node.h // 2)
        draw_arrow(draw, start, end)

    canvas.save(path)


def ensure_flowcharts() -> None:
    overview_nodes = [
        FlowNode("上传 doc/docx 素材", 120, 170, fill="#F4F8FB"),
        FlowNode("解析正文 / 表格 / 图片", 420, 170, fill="#F4F8FB"),
        FlowNode("生成目录与页摘要", 720, 170, fill="#F4F8FB"),
        FlowNode("用户确认并修改", 1020, 170, fill="#F4F8FB"),
        FlowNode("按固定模板导出 PPT", 420, 430, 500, 90, fill="#EEF7E8", outline="#8CC121"),
    ]
    overview_arrows = [(0, 1), (1, 2), (2, 3), (3, 4)]
    create_flowchart(FLOW_OVERVIEW_PATH, "业务流程图", overview_nodes, overview_arrows)

    export_nodes = [
        FlowNode("统一目录编号", 100, 180, fill="#F4F8FB"),
        FlowNode("统一页标题编号", 370, 180, fill="#F4F8FB"),
        FlowNode("按槽位控字数", 640, 180, fill="#F4F8FB"),
        FlowNode("映射模板页型", 910, 180, fill="#F4F8FB"),
        FlowNode("替换文本和图片关系", 370, 450, fill="#FDEEEE", outline="#C00000"),
        FlowNode("输出可编辑 .pptx", 760, 450, fill="#EEF7E8", outline="#8CC121"),
    ]
    export_arrows = [(0, 1), (1, 2), (2, 3), (1, 4), (3, 4), (4, 5)]
    create_flowchart(FLOW_EXPORT_PATH, "导出引擎流程图", export_nodes, export_arrows)


def set_page_layout(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.0)
    section.header_distance = Cm(1.0)
    section.footer_distance = Cm(1.0)


def set_document_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)

    for style_name in ("Title", "Heading 1", "Heading 2", "Heading 3"):
        style = document.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def add_title_page(document: Document) -> None:
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("固定模板 PPT 生成 Demo\nPRD")
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(24)
    run.bold = True

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("根据当前项目实现反向整理")
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(12)

    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run("更新时间：2026-03-10")
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(10.5)

    document.add_page_break()


def add_flowchart_section(document: Document) -> None:
    heading = document.add_heading("流程图", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    paragraph = document.add_paragraph("以下流程图用于说明当前 Demo 的业务闭环和导出引擎逻辑。")
    paragraph.paragraph_format.space_after = Pt(6)

    document.add_picture(str(FLOW_OVERVIEW_PATH), width=Inches(6.4))
    cap1 = document.add_paragraph("图 1 业务流程图")
    cap1.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_picture(str(FLOW_EXPORT_PATH), width=Inches(6.4))
    cap2 = document.add_paragraph("图 2 导出引擎流程图")
    cap2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_page_break()


def add_markdown_content(document: Document, markdown: str) -> None:
    for raw_line in markdown.splitlines():
        line = raw_line.rstrip()
        if not line:
            continue
        if line.startswith("# "):
            document.add_heading(line[2:].strip(), level=1)
            continue
        if line.startswith("## "):
            document.add_heading(line[3:].strip(), level=2)
            continue
        if line.startswith("### "):
            document.add_heading(line[4:].strip(), level=3)
            continue
        if line.startswith("- "):
            document.add_paragraph(line[2:].strip(), style="List Bullet")
            continue
        if line[:3].isdigit() and line[1:3] == ". ":
            document.add_paragraph(line[3:].strip(), style="List Number")
            continue
        document.add_paragraph(line)


def add_footer(document: Document) -> None:
    section = document.sections[0]
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("固定模板 PPT 生成 Demo PRD")
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(9)


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    TMP_DIR.mkdir(parents=True, exist_ok=True)
    ensure_flowcharts()

    markdown = MARKDOWN_PATH.read_text(encoding="utf-8")
    document = Document()
    set_page_layout(document)
    set_document_styles(document)
    add_title_page(document)
    add_flowchart_section(document)
    add_markdown_content(document, markdown)
    add_footer(document)
    document.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    main()
