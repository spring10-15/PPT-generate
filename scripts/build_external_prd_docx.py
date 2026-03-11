from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

from build_prd_docx import FLOW_EXPORT_PATH, FLOW_OVERVIEW_PATH, ensure_flowcharts


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "output" / "doc"
OUTPUT_PATH = OUTPUT_DIR / "固定模板PPT生成Demo-对外PRD.docx"

PRIMARY = "0070BF"
GREEN = "8CC121"
RED = "C00000"
LIGHT = "EEF4F9"
TEXT = "17324D"
MUTED = "5D738A"
BORDER = "D7E2EB"


DOCUMENT_INFO = [
    ("项目名称", "固定模板 PPT 生成 Demo"),
    ("文档类型", "产品需求文档（PRD）"),
    ("文档版本", "V1.2"),
    ("文档状态", "可对外评审"),
    ("更新日期", "2026-03-10"),
]

REVISION_HISTORY = [
    ("V0.9", "2026-03-10", "内部方案收敛", "形成固定模板导出、确认页与流程说明"),
    ("V1.0", "2026-03-10", "对外版整理", "补充范围、流程图、需求表与验收标准"),
    ("V1.1", "2026-03-10", "正式交付增强", "补充成功标准、里程碑、角色职责与非功能要求"),
    ("V1.2", "2026-03-10", "对外版完善", "补充交付边界、后续规划，并完善章节结构"),
]

PROJECT_SUMMARY = [
    "本项目面向需要快速生成标准化汇报材料的企业用户，目标是在固定 PPT 模板不变的前提下，实现从 Word 素材到 PowerPoint 的自动化生成。",
    "系统自动完成目录规划、详情摘要提炼、模板槽位填充和可编辑 .pptx 导出，用户只需做少量确认，不再手工排版。",
    "产品重点是模板保真、内容控字数和流程闭环，而不是自由设计模板或开放式生成任意样式幻灯片。",
]

SUCCESS_ROWS = [
    ("S-01", "生成效率", "相较人工排版显著缩短首版输出时间", "用户上传材料后可在一个流程内完成初版 PPT 生成"),
    ("S-02", "模板一致性", "输出 PPT 与固定模板在版式和样式层面保持一致", "不出现明显脱模、错位或模板提示词残留"),
    ("S-03", "可编辑性", "导出文件可继续在 PowerPoint 中编辑", "标题、正文、表格和图片保持为可编辑对象"),
    ("S-04", "人工校正成本", "用户只需做少量目录和摘要调整", "确认界面支持编辑和删除，不需要重新排版"),
]

SCOPE_ROWS = [
    ("范围内", "上传 doc/docx 素材", "解析正文、标题、表格与图片"),
    ("范围内", "生成目录架构", "目录优先控制在 4 项以内，可跨多页"),
    ("范围内", "生成每页摘要", "支持网页内编辑、删除、重排页码"),
    ("范围内", "固定模板导出", "输出可编辑 .pptx，并保留模板样式"),
    ("范围外", "用户自定义模板", "当前版本只支持固定模板"),
    ("范围外", "自动联网找图", "图片仅使用原文档内素材"),
    ("范围外", "AI 生成图片", "当前不纳入第一版"),
    ("范围外", "任意文档格式", "当前仅支持 doc/docx"),
]

INPUT_OUTPUT_ROWS = [
    ("输入", "源文件", "支持 doc/docx"),
    ("输入", "汇报人名称", "用于封面署名"),
    ("输入", "页数", "总页数，含封面和目录"),
    ("确认", "目录架构", "可修改标题、删除目录项"),
    ("确认", "详情摘要", "可修改页标题、摘要并删除详情页"),
    ("输出", "PPT 成品", "固定模板、可编辑 .pptx"),
]

FUNCTION_ROWS = [
    ("FR-01", "文档解析", "解析正文、标题、表格、图片；doc 在有 LibreOffice 时优先转 docx 解析", "上传示例文档后可得到结构化内容"),
    ("FR-02", "目录规划", "目录优先不超过 4 项，不同目录项不能合并到同一页", "生成结果目录层级清晰且页数可落地"),
    ("FR-03", "标题统一", "目录使用【一】【二】一级标题；详情页内容简介使用【1】【2】二级标题", "导出结果无混合编号格式"),
    ("FR-04", "确认交互", "目录和详情摘要分两个 tab 展示，均支持编辑和删除", "删除后页码自动重算"),
    ("FR-05", "模板复用", "详情页可重复使用模板页型，不允许自由改变版式", "不同页型可复用但位置和样式不变"),
    ("FR-06", "导出保真", "直接复用模板页 XML 槽位填充，不通过代码近似重画版式", "导出后与模板版式一致"),
    ("FR-07", "控字数", "按不同槽位容量裁剪与重组内容，避免挤爆模板文本框", "长文本页不出现明显模板残留说明文字"),
    ("FR-08", "图片处理", "只使用原文档图片，没有图片则保留图片区空位", "不引入外部图片"),
]

NON_FUNCTION_ROWS = [
    ("NFR-01", "模板保真", "不得改变固定模板中的位置、字号、颜色和行距规则"),
    ("NFR-02", "稳定性", "模型失败时仍能回退到规则化目录和摘要生成"),
    ("NFR-03", "可维护性", "模板页映射和槽位填充逻辑需要可扩展、可重复生成"),
    ("NFR-04", "可审阅性", "用户在导出前必须能看到目录和摘要的确认结果"),
    ("NFR-05", "性能预期", "单份普通 Word 文档应在可接受时间内完成初版生成"),
]

TEMPLATE_ROWS = [
    ("封面", "标题 / 副标题 / 署名 / 日期", "按模板第 1 页原位置填充"),
    ("目录页", "一级目录与页码范围", "按模板风格生成，不改变固定配色"),
    ("常规标题", "14pt / 加粗 / 白色", "沿用模板文本框样式"),
    ("强调标题", "14pt / 加粗 / 黄色", "沿用模板文本框样式"),
    ("小标题", "14pt / 模板红色强调", "沿用模板文本框样式"),
    ("说明性文字", "12pt / 模板灰色", "沿用模板文本框样式"),
    ("详情页版型", "4 / 5 / 9 / 10 / 11 / 12 / 13 / 15 / 16 页", "按内容类型自动匹配并可重复使用"),
]

MILESTONE_ROWS = [
    ("M1", "素材解析打通", "完成 doc/docx 上传、正文/表格/图片解析", "已完成"),
    ("M2", "目录与摘要确认", "完成目录生成、摘要生成、网页端确认交互", "已完成"),
    ("M3", "模板保真导出", "完成固定模板复用、槽位填充和 PPT 导出", "已完成"),
    ("M4", "复杂页型深化", "继续细化复杂页型槽位利用率和内容承载能力", "建议下一阶段"),
]

ROLE_ROWS = [
    ("产品", "定义流程、输入输出边界、验收标准和迭代优先级"),
    ("设计/模板维护", "维护固定模板样式、页面规范和模板变更基线"),
    ("研发", "实现解析、规划、确认页交互和模板导出链路"),
    ("业务用户", "上传素材、确认目录与摘要、校验输出结果"),
]

DELIVERY_ROWS = [
    ("交付内容", "网页 Demo、固定模板导出能力、PRD 文档、生成脚本"),
    ("不含内容", "用户自定义模板、多模板市场、自动找图、AI 生图"),
    ("依赖条件", "固定模板稳定、用户提供 doc/docx 素材、模型接口可用"),
    ("后续空间", "更多详情页模板覆盖、复杂图表页支持、版式校验自动化"),
]

ACCEPTANCE_ROWS = [
    ("AC-01", "用户可上传 doc/docx 并生成目录与摘要", "通过"),
    ("AC-02", "确认区支持修改目录标题、修改页标题、修改摘要", "通过"),
    ("AC-03", "确认区支持删除目录项与详情页，且页码自动重排", "通过"),
    ("AC-04", "导出后目录标题与页标题使用统一编号风格", "通过"),
    ("AC-05", "导出后不残留参考版式、内容简介、添加标题等模板示例词", "通过"),
    ("AC-06", "导出文件为可编辑 .pptx，且可正常打开", "通过"),
]

RISK_ROWS = [
    ("R-01", "doc 解析质量受环境依赖影响", "缺少 LibreOffice 时退化为文本级解析"),
    ("R-02", "模板更换会带来 XML 槽位映射重校准成本", "需重新做模板页槽位标定"),
    ("R-03", "复杂版式的可用容量差异大", "需要持续迭代控字数和槽位分配"),
    ("R-04", "自动目录规划可能仍需用户微调", "保留确认与编辑环节"),
]

NEXT_PHASE_ROWS = [
    ("P1", "扩展更多模板详情页类型", "提高复杂材料的内容利用率"),
    ("P2", "增加版式回归校验", "自动检测模板残留、越界和空白槽位"),
    ("P3", "提升 doc 图片解析能力", "降低 legacy 文档的解析退化"),
    ("P4", "支持更细粒度摘要编辑", "例如目录项级别批量改写、页内重排"),
]


def set_run_font(run, size: float, bold: bool = False, color: str = TEXT) -> None:
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def set_table_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, *, bold: bool = False, color: str = TEXT, size: float = 10.5) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    set_run_font(run, size, bold=bold, color=color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_doc_style(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.3)
    section.right_margin = Cm(2.0)
    section.header_distance = Cm(0.9)
    section.footer_distance = Cm(0.9)

    normal = document.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)

    for style_name, size in (("Title", 24), ("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11.5)):
        style = document.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(TEXT)


def add_field(paragraph, instruction: str) -> None:
    fld_simple = OxmlElement("w:fldSimple")
    fld_simple.set(qn("w:instr"), instruction)
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = ""
    run.append(text)
    fld_simple.append(run)
    paragraph._p.append(fld_simple)


def add_cover(document: Document) -> None:
    banner = document.add_table(rows=1, cols=1)
    banner.alignment = WD_TABLE_ALIGNMENT.CENTER
    banner.autofit = True
    banner_cell = banner.cell(0, 0)
    set_table_cell_shading(banner_cell, PRIMARY)
    paragraph = banner_cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("固定模板 PPT 生成 Demo")
    set_run_font(run, 20, bold=True, color="FFFFFF")

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("产品需求文档（PRD）")
    set_run_font(run, 26, bold=True, color=PRIMARY)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("对外评审版")
    set_run_font(run, 13, color=MUTED)

    document.add_paragraph()
    info_table = document.add_table(rows=0, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_table.style = "Table Grid"
    for label, value in DOCUMENT_INFO:
        row = info_table.add_row().cells
        set_table_cell_shading(row[0], LIGHT)
        set_cell_text(row[0], label, bold=True, color=PRIMARY)
        set_cell_text(row[1], value)

    document.add_paragraph()
    note = document.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = note.add_run("本文件用于对外沟通产品范围、流程、核心能力与验收标准。")
    set_run_font(run, 10.5, color=MUTED)

    document.add_page_break()


def add_revision_history(document: Document) -> None:
    document.add_heading("文档信息与修订记录", level=1)
    document.add_paragraph("以下信息用于说明当前文档版本、状态及修订历史。")
    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["版本", "日期", "阶段", "变更说明"]
    for cell, text in zip(table.rows[0].cells, headers):
        set_table_cell_shading(cell, LIGHT)
        set_cell_text(cell, text, bold=True, color=PRIMARY)
    for row_data in REVISION_HISTORY:
        row = table.add_row().cells
        for cell, value in zip(row, row_data):
            set_cell_text(cell, value)


def add_toc(document: Document) -> None:
    document.add_page_break()
    document.add_heading("目录", level=1)
    paragraph = document.add_paragraph()
    add_field(paragraph, 'TOC \\o "1-3" \\h \\z \\u')
    tip = document.add_paragraph()
    run = tip.add_run("目录字段已插入，可在 Word 中更新后显示页码。")
    set_run_font(run, 9.5, color=MUTED)
    document.add_page_break()


def add_summary_section(document: Document) -> None:
    document.add_heading("1. 项目概览", level=1)
    document.add_heading("1.1 项目背景", level=2)
    for item in PROJECT_SUMMARY:
        document.add_paragraph(item, style="List Bullet")

    document.add_heading("1.2 成功标准", level=2)
    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell, text in zip(table.rows[0].cells, ["编号", "指标项", "目标", "判定口径"]):
        set_table_cell_shading(cell, LIGHT)
        set_cell_text(cell, text, bold=True, color=PRIMARY)
    for row_data in SUCCESS_ROWS:
        row = table.add_row().cells
        for cell, value in zip(row, row_data):
            set_cell_text(cell, value)

    document.add_heading("1.3 范围说明", level=2)
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell, text in zip(table.rows[0].cells, ["分类", "事项", "说明"]):
        set_table_cell_shading(cell, LIGHT)
        set_cell_text(cell, text, bold=True, color=PRIMARY)
    for row_data in SCOPE_ROWS:
        row = table.add_row().cells
        for cell, value in zip(row, row_data):
            set_cell_text(cell, value)


def add_flow_sections(document: Document) -> None:
    document.add_heading("2. 核心流程", level=1)
    paragraph = document.add_paragraph("系统流程分为用户业务流程和导出引擎流程两部分。")
    paragraph.paragraph_format.space_after = Pt(6)

    document.add_heading("2.1 业务流程图", level=2)
    document.add_picture(str(FLOW_OVERVIEW_PATH), width=Inches(6.3))
    cap1 = document.add_paragraph("图 1 业务流程图")
    cap1.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_heading("2.2 导出引擎流程图", level=2)
    document.add_picture(str(FLOW_EXPORT_PATH), width=Inches(6.3))
    cap2 = document.add_paragraph("图 2 导出引擎流程图")
    cap2.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_input_output(document: Document) -> None:
    document.add_heading("3. 输入输出与交互", level=1)
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell, text in zip(table.rows[0].cells, ["类型", "项目", "说明"]):
        set_table_cell_shading(cell, LIGHT)
        set_cell_text(cell, text, bold=True, color=PRIMARY)
    for row_data in INPUT_OUTPUT_ROWS:
        row = table.add_row().cells
        for cell, value in zip(row, row_data):
            set_cell_text(cell, value)

    document.add_heading("3.1 界面交互要求", level=2)
    for item in [
        "首页输入项仅保留：源文件、汇报人名称、页数。",
        "确认区使用两个 tab：目录架构、详情摘要。",
        "目录架构 tab 支持编辑标题、删除目录项。",
        "详情摘要 tab 支持编辑页标题、编辑摘要、删除详情页。",
        "删除目录项或详情页后，页码和总页数自动重排。",
    ]:
        document.add_paragraph(item, style="List Bullet")


def add_function_requirements(document: Document) -> None:
    document.add_heading("4. 功能需求", level=1)
    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["编号", "模块", "需求说明", "验收方式"]
    for cell, text in zip(table.rows[0].cells, headers):
        set_table_cell_shading(cell, LIGHT)
        set_cell_text(cell, text, bold=True, color=PRIMARY)
    for row_data in FUNCTION_ROWS:
        row = table.add_row().cells
        for cell, value in zip(row, row_data):
            set_cell_text(cell, value)


def add_non_function_requirements(document: Document) -> None:
    document.add_heading("5. 非功能要求", level=1)
    intro = document.add_paragraph(
        "以下要求用于约束对外版本在模板一致性、稳定性和可维护性上的最低基线。"
    )
    intro.paragraph_format.space_after = Pt(6)

    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell, text in zip(table.rows[0].cells, ["编号", "分类", "要求说明"]):
        set_table_cell_shading(cell, LIGHT)
        set_cell_text(cell, text, bold=True, color=PRIMARY)
    for row_data in NON_FUNCTION_ROWS:
        row = table.add_row().cells
        for cell, value in zip(row, row_data):
            set_cell_text(cell, value)


def add_template_fidelity(document: Document) -> None:
    document.add_heading("6. 模板保真要求", level=1)
    preface = document.add_paragraph(
        "本项目的核心是固定模板保真。系统必须复制模板中的真实文本槽位、图片区和结构元素，不允许通过近似绘制替代。"
    )
    preface.paragraph_format.space_after = Pt(6)

    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell, text in zip(table.rows[0].cells, ["对象", "要求", "说明"]):
        set_table_cell_shading(cell, LIGHT)
        set_cell_text(cell, text, bold=True, color=PRIMARY)
    for row_data in TEMPLATE_ROWS:
        row = table.add_row().cells
        for cell, value in zip(row, row_data):
            set_cell_text(cell, value)


def add_milestones(document: Document) -> None:
    document.add_heading("7. 实施里程碑", level=1)
    intro = document.add_paragraph("当前版本的落地节奏和后续工作重点如下。")
    intro.paragraph_format.space_after = Pt(6)

    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell, text in zip(table.rows[0].cells, ["里程碑", "阶段名称", "交付内容", "状态"]):
        set_table_cell_shading(cell, LIGHT)
        set_cell_text(cell, text, bold=True, color=PRIMARY)
    for row_data in MILESTONE_ROWS:
        row = table.add_row().cells
        for cell, value in zip(row, row_data):
            set_cell_text(cell, value)


def add_roles(document: Document) -> None:
    document.add_heading("8. 角色与职责", level=1)
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell, text in zip(table.rows[0].cells, ["角色", "职责说明"]):
        set_table_cell_shading(cell, LIGHT)
        set_cell_text(cell, text, bold=True, color=PRIMARY)
    for row_data in ROLE_ROWS:
        row = table.add_row().cells
        for cell, value in zip(row, row_data):
            set_cell_text(cell, value)


def add_delivery_boundaries(document: Document) -> None:
    document.add_heading("9. 交付边界", level=1)
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell, text in zip(table.rows[0].cells, ["项目项", "说明"]):
        set_table_cell_shading(cell, LIGHT)
        set_cell_text(cell, text, bold=True, color=PRIMARY)
    for row_data in DELIVERY_ROWS:
        row = table.add_row().cells
        for cell, value in zip(row, row_data):
            set_cell_text(cell, value)


def add_acceptance(document: Document) -> None:
    document.add_heading("10. 验收标准", level=1)
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell, text in zip(table.rows[0].cells, ["编号", "验收项", "状态"]):
        set_table_cell_shading(cell, LIGHT)
        set_cell_text(cell, text, bold=True, color=PRIMARY)
    for row_data in ACCEPTANCE_ROWS:
        row = table.add_row().cells
        for cell, value in zip(row, row_data):
            set_cell_text(cell, value)


def add_risks(document: Document) -> None:
    document.add_heading("11. 风险与边界", level=1)
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell, text in zip(table.rows[0].cells, ["编号", "风险描述", "当前处理方式"]):
        set_table_cell_shading(cell, LIGHT)
        set_cell_text(cell, text, bold=True, color=PRIMARY)
    for row_data in RISK_ROWS:
        row = table.add_row().cells
        for cell, value in zip(row, row_data):
            set_cell_text(cell, value)


def add_next_phase(document: Document) -> None:
    document.add_heading("12. 下一阶段建议", level=1)
    paragraph = document.add_paragraph(
        "在当前 Demo 已实现端到端流程和模板保真导出的基础上，建议按以下方向继续深化。"
    )
    paragraph.paragraph_format.space_after = Pt(6)

    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell, text in zip(table.rows[0].cells, ["编号", "建议事项", "目标"]):
        set_table_cell_shading(cell, LIGHT)
        set_cell_text(cell, text, bold=True, color=PRIMARY)
    for row_data in NEXT_PHASE_ROWS:
        row = table.add_row().cells
        for cell, value in zip(row, row_data):
            set_cell_text(cell, value)


def add_header_footer(document: Document) -> None:
    for section in document.sections:
        section.start_type = WD_SECTION_START.NEW_PAGE
        header = section.header.paragraphs[0]
        header.alignment = WD_ALIGN_PARAGRAPH.LEFT
        header_run = header.add_run("固定模板 PPT 生成 Demo | 产品需求文档（对外版）")
        set_run_font(header_run, 9, color=MUTED)

        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer.add_run("第 ")
        set_run_font(footer_run, 9, color=MUTED)
        add_field(footer, "PAGE")
        footer_run2 = footer.add_run(" 页")
        set_run_font(footer_run2, 9, color=MUTED)


def build_document() -> Document:
    ensure_flowcharts()
    document = Document()
    set_doc_style(document)
    add_cover(document)
    add_revision_history(document)
    add_toc(document)
    add_summary_section(document)
    add_flow_sections(document)
    add_input_output(document)
    add_function_requirements(document)
    add_non_function_requirements(document)
    add_template_fidelity(document)
    add_milestones(document)
    add_roles(document)
    add_delivery_boundaries(document)
    add_acceptance(document)
    add_risks(document)
    add_next_phase(document)
    add_header_footer(document)
    return document


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    document = build_document()
    document.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    main()
